"""
Standard professional resume template.

Deliberately plain and ATS-friendly: single column, no tables/text boxes,
standard heading styles, Calibri throughout. This is the deterministic
rendering step — it never generates wording, it only lays out structured
data that the tailoring stage (Claude) has already produced.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

FONT_NAME = "Calibri"
ACCENT_COLOR = RGBColor(0x1F, 0x2A, 0x44)  # dark navy, professional not flashy
BODY_SIZE = Pt(10.5)
NAME_SIZE = Pt(20)
SECTION_SIZE = Pt(12)


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    # Ensure east-asian font fallback doesn't override on some systems
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.size = SECTION_SIZE
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR
    # simple bottom border for a clean section divider
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(
        qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "1F2A44"},
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = BODY_SIZE


def _format_date(iso_date: str | None) -> str:
    if not iso_date:
        return "Present"
    # Expecting "YYYY-MM-DD" or "YYYY-MM" from the web app; keep this loose
    # since the web layer owns real date parsing/validation.
    parts = iso_date.split("-")
    if len(parts) >= 2:
        return f"{parts[1]}/{parts[0]}"
    return iso_date


def build_resume_docx(data: dict[str, Any]) -> BytesIO:
    """
    Expected shape (produced by the web app from tailored profile data):

    {
      "fullName": str,
      "headline": str | None,
      "contact": {"email": str, "phone": str|None, "location": str|None,
                   "linkedin": str|None, "github": str|None},
      "summary": str | None,
      "experiences": [
        {"organization": str, "title": str, "location": str|None,
         "startDate": str, "endDate": str|None, "bullets": [str, ...]}
      ],
      "skills": [str, ...],
      "education": [
        {"institution": str, "degree": str, "field": str|None,
         "endDate": str|None}
      ]
    }
    """
    doc = Document()
    _set_default_font(doc)

    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(50)
        section.right_margin = Pt(50)

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data["fullName"])
    name_run.font.size = NAME_SIZE
    name_run.font.bold = True
    name_run.font.color.rgb = ACCENT_COLOR

    # Headline
    if data.get("headline"):
        h_p = doc.add_paragraph()
        h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_run = h_p.add_run(data["headline"])
        h_run.font.size = Pt(11)
        h_run.italic = True

    # Contact line
    contact = data.get("contact", {})
    contact_parts = [
        v for v in [
            contact.get("email"),
            contact.get("phone"),
            contact.get("location"),
            contact.get("linkedin"),
            contact.get("github"),
        ] if v
    ]
    if contact_parts:
        c_p = doc.add_paragraph()
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = c_p.add_run(" | ".join(contact_parts))
        c_run.font.size = Pt(9.5)

    # Summary
    if data.get("summary"):
        _add_section_heading(doc, "Summary")
        s_p = doc.add_paragraph()
        s_run = s_p.add_run(data["summary"])
        s_run.font.size = BODY_SIZE

    # Experience
    experiences = data.get("experiences", [])
    if experiences:
        _add_section_heading(doc, "Experience")
        for exp in experiences:
            role_p = doc.add_paragraph()
            role_p.paragraph_format.space_after = Pt(0)
            title_run = role_p.add_run(f"{exp['title']} — {exp['organization']}")
            title_run.bold = True
            title_run.font.size = BODY_SIZE

            date_range = f"{_format_date(exp.get('startDate'))} – {_format_date(exp.get('endDate'))}"
            meta_bits = [date_range]
            if exp.get("location"):
                meta_bits.append(exp["location"])
            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.space_after = Pt(2)
            meta_run = meta_p.add_run(" | ".join(meta_bits))
            meta_run.italic = True
            meta_run.font.size = Pt(9.5)

            for bullet in exp.get("bullets", []):
                _add_bullet(doc, bullet)

    # Skills
    skills = data.get("skills", [])
    if skills:
        _add_section_heading(doc, "Skills")
        skills_p = doc.add_paragraph()
        skills_run = skills_p.add_run(" • ".join(skills))
        skills_run.font.size = BODY_SIZE

    # Education
    education = data.get("education", [])
    if education:
        _add_section_heading(doc, "Education")
        for edu in education:
            edu_p = doc.add_paragraph()
            edu_p.paragraph_format.space_after = Pt(2)
            degree_line = edu["degree"]
            if edu.get("field"):
                degree_line += f", {edu['field']}"
            edu_run = edu_p.add_run(f"{degree_line} — {edu['institution']}")
            edu_run.bold = True
            edu_run.font.size = BODY_SIZE
            if edu.get("endDate"):
                date_run = edu_p.add_run(f"  ({_format_date(edu['endDate'])})")
                date_run.italic = True
                date_run.font.size = Pt(9.5)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
